#!/usr/bin/env python3
"""
adtool.py — deterministic AD builder (CODIFIED mode) + Risk support

Spec directory files (YAML):
  - stakeholders.yaml              (required)
  - concerns.yaml                  (required)
  - capabilities.yaml              (required)
  - service_levels.yaml            (optional)
  - risks.yaml                     (optional, but supported and rendered if present)

What it does (deterministic):
  1) Validation: basic structure, IDs, reference integrity
  2) Gap analysis: missing fields, uncovered links, conflicts, coverage checks
  3) Rendering: builds AD markdown from a stable template, inserts <TODO> markers
  4) Outputs:
     - AD.md (or provided --out)
     - gaps.md (TODO/gaps report)
     - validation_report.json (machine-readable report)

Dependencies:
  pip install pyyaml jinja2

Usage:
  python adtool.py validate ad/spec --report ad/output/validation_report.json
  python adtool.py build ad/spec --out ad/output/AD_codified.md --format md
  python adtool.py build ad/spec --out ad/output/AD_codified.docx --format docx
"""

from __future__ import annotations

import argparse
import dataclasses
import datetime as dt
import json
import re
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import yaml
from jinja2 import Environment, BaseLoader, StrictUndefined


# ---------------------------
# Constants / helpers
# ---------------------------

TODO_HTML = '<span style="color:red"><b><TODO></b></span>'
TODO_CONFLICT_HTML = '<span style="color:red"><b><TODO-CONFLICT></b></span>'

DEFAULT_FILES = {
    "stakeholders": ("stakeholders.yaml", "stakeholders.yml"),
    "concerns": ("concerns.yaml", "concerns.yml"),
    "capabilities": ("capabilities.yaml", "capabilities.yml"),
    "service_levels": ("service_levels.yaml", "service_levels.yml"),
    "risks": ("risks.yaml", "risks.yml"),
}

SEVERITY_ORDER = {"ERROR": 0, "WARN": 1, "INFO": 2}

# conservative but practical: STK-OPS, C-003, CAP-05, SL-001, R-001 etc.
ID_RE = re.compile(r"^[A-Za-z][A-Za-z0-9\-_.:]*$")


def todo(label: str = "TODO") -> str:
    if label == "TODO":
        return TODO_HTML
    if label.upper().startswith("TODO-CONFLICT"):
        return TODO_CONFLICT_HTML
    return f'<span style="color:red"><b><{label}></b></span>'


def now_iso() -> str:
    return dt.datetime.now(dt.timezone.utc).isoformat()


def as_one_line(s: Any) -> str:
    if s is None:
        return ""
    if not isinstance(s, str):
        s = str(s)
    return " ".join(s.replace("\r", "").replace("\n", " ").split()).strip()


# ---------------------------
# Data model
# ---------------------------

@dataclasses.dataclass(frozen=True)
class Stakeholder:
    id: str
    name: str
    description: str = ""


@dataclasses.dataclass(frozen=True)
class ConcernMeasurement:
    sli: str = ""
    slo: str = ""
    sla: str = ""
    service_level_id: str = ""


@dataclasses.dataclass(frozen=True)
class Concern:
    id: str
    name: str
    description: str
    stakeholders: List[str]
    tags: List[str]
    measurement: ConcernMeasurement


@dataclasses.dataclass(frozen=True)
class Capability:
    id: str
    name: str
    description: str
    addresses_concerns: List[str]
    constraints: Dict[str, Any]
    tags: List[str]


@dataclasses.dataclass(frozen=True)
class ServiceLevel:
    id: str
    name: str
    sli_definition: str
    window: str
    exclusions: str
    target_slo: str
    contractual_sla: str


@dataclasses.dataclass(frozen=True)
class Risk:
    id: str
    title: str
    description: str
    type: str               # e.g., Programmatic / Operational / Data / Security ...
    status: str             # e.g., Open / Mitigating / Closed
    owner: str              # stakeholder id
    affected_concerns: List[str]
    affected_capabilities: List[str]
    threatened_service_levels: List[str]
    linked_views: List[str] # e.g., ["AV-1", "AcV-2:M3", "SV-1"]
    mitigation: str         # free text or references to CAP IDs/actions


@dataclasses.dataclass
class Issue:
    severity: str     # ERROR / WARN / INFO
    code: str         # e.g., MISSING_FIELD, BROKEN_LINK, GAP, CONFLICT
    location: str     # e.g., "risks:R-001.owner"
    message: str

    def as_dict(self) -> Dict[str, Any]:
        return {
            "severity": self.severity,
            "code": self.code,
            "location": self.location,
            "message": self.message,
        }


# ---------------------------
# Spec loading
# ---------------------------

def load_yaml_file(path: Path) -> Dict[str, Any]:
    with path.open("r", encoding="utf-8") as f:
        data = yaml.safe_load(f)
    if data is None:
        return {}
    if not isinstance(data, dict):
        raise ValueError(f"YAML root must be a mapping (dict): {path}")
    return data


def find_file(spec_dir: Path, candidates: Tuple[str, ...]) -> Optional[Path]:
    for name in candidates:
        p = spec_dir / name
        if p.exists() and p.is_file():
            return p
    return None


def load_spec(spec_dir: Path) -> Tuple[Dict[str, Any], List[Issue]]:
    issues: List[Issue] = []
    spec_dir = spec_dir.resolve()

    loaded: Dict[str, Any] = {}

    for key, candidates in DEFAULT_FILES.items():
        fp = find_file(spec_dir, candidates)
        if fp is None:
            if key in ("service_levels", "risks"):
                loaded[key] = {}
                continue
            issues.append(Issue(
                severity="ERROR",
                code="MISSING_FILE",
                location=f"spec:{key}",
                message=f"Required file is missing in {spec_dir}: one of {list(candidates)}",
            ))
            loaded[key] = {}
            continue

        try:
            loaded[key] = load_yaml_file(fp)
        except Exception as e:
            issues.append(Issue(
                severity="ERROR",
                code="YAML_PARSE_ERROR",
                location=f"spec:{fp.name}",
                message=str(e),
            ))
            loaded[key] = {}

    return loaded, issues


# ---------------------------
# Validation & parsing helpers
# ---------------------------

def _require_str(d: Dict[str, Any], field: str, location: str, issues: List[Issue], required: bool = True) -> str:
    v = d.get(field, "")
    if v is None:
        v = ""
    if not isinstance(v, str):
        issues.append(Issue("ERROR", "TYPE_ERROR", f"{location}.{field}", f"Expected string, got {type(v).__name__}"))
        return ""
    v = v.strip()
    if required and not v:
        issues.append(Issue("WARN", "MISSING_FIELD", f"{location}.{field}", "Missing/empty field"))
    return v


def _require_list_of_str(d: Dict[str, Any], field: str, location: str, issues: List[Issue], required: bool = False) -> List[str]:
    v = d.get(field, [])
    if v is None:
        v = []
    if not isinstance(v, list):
        issues.append(Issue("ERROR", "TYPE_ERROR", f"{location}.{field}", f"Expected list, got {type(v).__name__}"))
        return []
    out: List[str] = []
    for i, item in enumerate(v):
        if not isinstance(item, str):
            issues.append(Issue("ERROR", "TYPE_ERROR", f"{location}.{field}[{i}]", f"Expected string, got {type(item).__name__}"))
            continue
        s = item.strip()
        if s:
            out.append(s)
    if required and not out:
        issues.append(Issue("WARN", "MISSING_FIELD", f"{location}.{field}", "Missing/empty list"))
    return out


def _require_dict(d: Dict[str, Any], field: str, location: str, issues: List[Issue]) -> Dict[str, Any]:
    v = d.get(field, {})
    if v is None:
        v = {}
    if not isinstance(v, dict):
        issues.append(Issue("ERROR", "TYPE_ERROR", f"{location}.{field}", f"Expected dict, got {type(v).__name__}"))
        return {}
    return v


def _validate_id(id_value: str, location: str, issues: List[Issue]) -> None:
    if not id_value:
        issues.append(Issue("ERROR", "MISSING_ID", location, "Missing id"))
        return
    if not ID_RE.match(id_value):
        issues.append(Issue("ERROR", "BAD_ID_FORMAT", location, f"ID '{id_value}' has invalid format"))


# ---------------------------
# Parse entities
# ---------------------------

def parse_stakeholders(raw: Dict[str, Any], issues: List[Issue]) -> List[Stakeholder]:
    items = raw.get("stakeholders", []) or []
    if not isinstance(items, list):
        issues.append(Issue("ERROR", "TYPE_ERROR", "stakeholders", "Expected list at key 'stakeholders'"))
        return []

    out: List[Stakeholder] = []
    seen: set[str] = set()

    for idx, it in enumerate(items):
        loc = f"stakeholders[{idx}]"
        if not isinstance(it, dict):
            issues.append(Issue("ERROR", "TYPE_ERROR", loc, f"Expected dict, got {type(it).__name__}"))
            continue

        sid = _require_str(it, "id", loc, issues, required=True)
        _validate_id(sid, f"{loc}.id", issues)
        if sid in seen:
            issues.append(Issue("ERROR", "DUPLICATE_ID", f"{loc}.id", f"Duplicate stakeholder id '{sid}'"))
        seen.add(sid)

        name = _require_str(it, "name", loc, issues, required=True)
        desc = _require_str(it, "description", loc, issues, required=False)

        out.append(Stakeholder(id=sid, name=name, description=desc))

    out.sort(key=lambda x: x.id)
    return out


def parse_concerns(raw: Dict[str, Any], issues: List[Issue]) -> List[Concern]:
    items = raw.get("concerns", []) or []
    if not isinstance(items, list):
        issues.append(Issue("ERROR", "TYPE_ERROR", "concerns", "Expected list at key 'concerns'"))
        return []

    out: List[Concern] = []
    seen: set[str] = set()

    for idx, it in enumerate(items):
        loc = f"concerns[{idx}]"
        if not isinstance(it, dict):
            issues.append(Issue("ERROR", "TYPE_ERROR", loc, f"Expected dict, got {type(it).__name__}"))
            continue

        cid = _require_str(it, "id", loc, issues, required=True)
        _validate_id(cid, f"{loc}.id", issues)
        if cid in seen:
            issues.append(Issue("ERROR", "DUPLICATE_ID", f"{loc}.id", f"Duplicate concern id '{cid}'"))
        seen.add(cid)

        name = _require_str(it, "name", loc, issues, required=True)
        desc = _require_str(it, "description", loc, issues, required=True)
        stks = _require_list_of_str(it, "stakeholders", loc, issues, required=True)
        tags = _require_list_of_str(it, "tags", loc, issues, required=False)

        meas_raw = it.get("measurement", {}) or {}
        if not isinstance(meas_raw, dict):
            issues.append(Issue("ERROR", "TYPE_ERROR", f"{loc}.measurement", f"Expected dict, got {type(meas_raw).__name__}"))
            meas_raw = {}

        meas = ConcernMeasurement(
            sli=_require_str(meas_raw, "sli", f"{loc}.measurement", issues, required=False),
            slo=_require_str(meas_raw, "slo", f"{loc}.measurement", issues, required=False),
            sla=_require_str(meas_raw, "sla", f"{loc}.measurement", issues, required=False),
            service_level_id=_require_str(meas_raw, "service_level_id", f"{loc}.measurement", issues, required=False),
        )

        out.append(Concern(
            id=cid,
            name=name,
            description=desc,
            stakeholders=stks,
            tags=tags,
            measurement=meas,
        ))

    out.sort(key=lambda x: x.id)
    return out


def parse_capabilities(raw: Dict[str, Any], issues: List[Issue]) -> List[Capability]:
    items = raw.get("capabilities", []) or []
    if not isinstance(items, list):
        issues.append(Issue("ERROR", "TYPE_ERROR", "capabilities", "Expected list at key 'capabilities'"))
        return []

    out: List[Capability] = []
    seen: set[str] = set()

    for idx, it in enumerate(items):
        loc = f"capabilities[{idx}]"
        if not isinstance(it, dict):
            issues.append(Issue("ERROR", "TYPE_ERROR", loc, f"Expected dict, got {type(it).__name__}"))
            continue

        capid = _require_str(it, "id", loc, issues, required=True)
        _validate_id(capid, f"{loc}.id", issues)
        if capid in seen:
            issues.append(Issue("ERROR", "DUPLICATE_ID", f"{loc}.id", f"Duplicate capability id '{capid}'"))
        seen.add(capid)

        name = _require_str(it, "name", loc, issues, required=True)
        desc = _require_str(it, "description", loc, issues, required=True)
        concerns = _require_list_of_str(it, "addresses_concerns", loc, issues, required=False)
        if not concerns:
            issues.append(Issue("WARN", "GAP", f"{loc}.addresses_concerns", "Capability does not address any concerns"))

        constraints = _require_dict(it, "constraints", loc, issues)
        tags = _require_list_of_str(it, "tags", loc, issues, required=False)

        out.append(Capability(
            id=capid,
            name=name,
            description=desc,
            addresses_concerns=concerns,
            constraints=constraints,
            tags=tags,
        ))

    out.sort(key=lambda x: x.id)
    return out


def parse_service_levels(raw: Dict[str, Any], issues: List[Issue]) -> List[ServiceLevel]:
    items = raw.get("service_levels", []) or []
    if not isinstance(items, list):
        if raw:
            issues.append(Issue("ERROR", "TYPE_ERROR", "service_levels", "Expected list at key 'service_levels'"))
        return []

    out: List[ServiceLevel] = []
    seen: set[str] = set()

    for idx, it in enumerate(items):
        loc = f"service_levels[{idx}]"
        if not isinstance(it, dict):
            issues.append(Issue("ERROR", "TYPE_ERROR", loc, f"Expected dict, got {type(it).__name__}"))
            continue

        sid = _require_str(it, "id", loc, issues, required=True)
        _validate_id(sid, f"{loc}.id", issues)
        if sid in seen:
            issues.append(Issue("ERROR", "DUPLICATE_ID", f"{loc}.id", f"Duplicate service level id '{sid}'"))
        seen.add(sid)

        sl = ServiceLevel(
            id=sid,
            name=_require_str(it, "name", loc, issues, required=True),
            sli_definition=_require_str(it, "sli_definition", loc, issues, required=True),
            window=_require_str(it, "window", loc, issues, required=True),
            exclusions=_require_str(it, "exclusions", loc, issues, required=False),
            target_slo=_require_str(it, "target_slo", loc, issues, required=False),
            contractual_sla=_require_str(it, "contractual_sla", loc, issues, required=False),
        )
        out.append(sl)

    out.sort(key=lambda x: x.id)
    return out


def parse_risks(raw: Dict[str, Any], issues: List[Issue]) -> List[Risk]:
    items = raw.get("risks", []) or []
    if not isinstance(items, list):
        if raw:
            issues.append(Issue("ERROR", "TYPE_ERROR", "risks", "Expected list at key 'risks'"))
        return []

    out: List[Risk] = []
    seen: set[str] = set()

    for idx, it in enumerate(items):
        loc = f"risks[{idx}]"
        if not isinstance(it, dict):
            issues.append(Issue("ERROR", "TYPE_ERROR", loc, f"Expected dict, got {type(it).__name__}"))
            continue

        rid = _require_str(it, "id", loc, issues, required=True)
        _validate_id(rid, f"{loc}.id", issues)
        if rid in seen:
            issues.append(Issue("ERROR", "DUPLICATE_ID", f"{loc}.id", f"Duplicate risk id '{rid}'"))
        seen.add(rid)

        title = _require_str(it, "title", loc, issues, required=True)
        desc = _require_str(it, "description", loc, issues, required=False)
        rtype = _require_str(it, "type", loc, issues, required=True)
        status = _require_str(it, "status", loc, issues, required=True)
        owner = _require_str(it, "owner", loc, issues, required=True)

        affected_concerns = _require_list_of_str(it, "affected_concerns", loc, issues, required=False)
        affected_caps = _require_list_of_str(it, "affected_capabilities", loc, issues, required=False)
        threatened_sls = _require_list_of_str(it, "threatened_service_levels", loc, issues, required=False)
        linked_views = _require_list_of_str(it, "linked_views", loc, issues, required=False)
        mitigation = _require_str(it, "mitigation", loc, issues, required=False)

        out.append(Risk(
            id=rid,
            title=title,
            description=desc,
            type=rtype,
            status=status,
            owner=owner,
            affected_concerns=affected_concerns,
            affected_capabilities=affected_caps,
            threatened_service_levels=threatened_sls,
            linked_views=linked_views,
            mitigation=mitigation,
        ))

    out.sort(key=lambda x: x.id)
    return out


# ---------------------------
# Analysis (gaps, coverage, consistency)
# ---------------------------

def keyword_classify_capability(cap: Capability) -> List[str]:
    """
    Deterministic fallback classifier (only used if cap.tags is empty).
    Adds 'Operational' tag if name/description contains operational keywords.
    """
    text = f"{cap.name} {cap.description}".lower()
    operational_keywords = [
        "observ", "monitor", "metric", "log", "trace", "diagnos", "alert",
        "deploy", "release", "rollback", "migration", "runbook", "sre", "ops",
        "availability", "resilien", "incident",
    ]
    tags: List[str] = []
    if any(k in text for k in operational_keywords):
        tags.append("Operational")
    return tags


def is_programmatic_risk(risk_type: str) -> bool:
    """
    Heuristic: if type indicates program/acquisition/schedule, we expect AcV-2 link.
    Deterministic and configurable later if needed.
    """
    t = (risk_type or "").strip().lower()
    return any(k in t for k in ["program", "programme", "acquisition", "schedule", "timeline"])


def analyze(
    stakeholders: List[Stakeholder],
    concerns: List[Concern],
    capabilities: List[Capability],
    service_levels: List[ServiceLevel],
    risks: List[Risk],
) -> List[Issue]:
    issues: List[Issue] = []

    stk_ids = {s.id for s in stakeholders}
    concern_ids = {c.id for c in concerns}
    sl_ids = {s.id for s in service_levels}
    cap_ids = {c.id for c in capabilities}

    # Broken link checks: concern -> stakeholders
    for c in concerns:
        for sid in c.stakeholders:
            if sid not in stk_ids:
                issues.append(Issue("ERROR", "BROKEN_LINK", f"concerns:{c.id}.stakeholders",
                                   f"Concern references unknown stakeholder id '{sid}'"))

    # Broken link checks: capability -> concerns
    for cap in capabilities:
        for cid in cap.addresses_concerns:
            if cid not in concern_ids:
                issues.append(Issue("ERROR", "BROKEN_LINK", f"capabilities:{cap.id}.addresses_concerns",
                                   f"Capability references unknown concern id '{cid}'"))

    # Broken link checks: concern.measurement.service_level_id
    for c in concerns:
        if c.measurement.service_level_id and c.measurement.service_level_id not in sl_ids:
            issues.append(Issue("ERROR", "BROKEN_LINK", f"concerns:{c.id}.measurement.service_level_id",
                               f"Concern references unknown service level id '{c.measurement.service_level_id}'"))

    # Coverage: stakeholder must have at least one concern
    concerns_by_stk: Dict[str, List[str]] = {sid: [] for sid in stk_ids}
    for c in concerns:
        for sid in c.stakeholders:
            if sid in concerns_by_stk:
                concerns_by_stk[sid].append(c.id)
    for sid, lst in sorted(concerns_by_stk.items(), key=lambda x: x[0]):
        if not lst:
            issues.append(Issue("WARN", "GAP", f"stakeholders:{sid}", "Stakeholder has no concerns mapped to it"))

    # Coverage: each concern should be addressed by at least one capability
    caps_by_concern: Dict[str, List[str]] = {cid: [] for cid in concern_ids}
    for cap in capabilities:
        for cid in cap.addresses_concerns:
            if cid in caps_by_concern:
                caps_by_concern[cid].append(cap.id)
    for cid, lst in sorted(caps_by_concern.items(), key=lambda x: x[0]):
        if not lst:
            issues.append(Issue("WARN", "GAP", f"concerns:{cid}", "Concern is not addressed by any capability"))

    # Coverage: each capability should address at least one concern
    for cap in capabilities:
        if not cap.addresses_concerns:
            issues.append(Issue("WARN", "GAP", f"capabilities:{cap.id}", "Capability is not linked to any concerns"))

    # Operational concerns should have some measurement (SLI or service_level_id)
    for c in concerns:
        tags = {t.strip() for t in c.tags}
        if "Operational" in tags:
            if not (c.measurement.sli or c.measurement.service_level_id):
                issues.append(Issue("WARN", "MISSING_MEASUREMENT", f"concerns:{c.id}.measurement",
                                   "Operational concern is missing measurement.sli or measurement.service_level_id"))

    # If SLO/SLA present but no SLI + no service_level_id => incomplete definition
    for c in concerns:
        if (c.measurement.slo or c.measurement.sla) and not (c.measurement.sli or c.measurement.service_level_id):
            issues.append(Issue("WARN", "INCOMPLETE_SLO_SLA", f"concerns:{c.id}.measurement",
                               "SLO/SLA stated but no SLI (or service_level_id) defined"))

    # Double-tag concerns (Business + Operational) should be addressed by at least one operational capability
    operational_caps: set[str] = set()
    for cap in capabilities:
        cap_tags = {t.strip() for t in cap.tags} if cap.tags else set(keyword_classify_capability(cap))
        if "Operational" in cap_tags:
            operational_caps.add(cap.id)

    for c in concerns:
        ctags = {t.strip() for t in c.tags}
        if "Business" in ctags and "Operational" in ctags:
            linked_caps = caps_by_concern.get(c.id, [])
            if linked_caps and not any(cap_id in operational_caps for cap_id in linked_caps):
                issues.append(Issue("WARN", "GAP", f"concerns:{c.id}",
                                   "Double-tag concern (Business+Operational) has no operational capability linked"))

    # Service level catalog: required fields
    for sl in service_levels:
        if not sl.sli_definition.strip():
            issues.append(Issue("WARN", "MISSING_FIELD", f"service_levels:{sl.id}.sli_definition", "Missing SLI definition"))
        if not sl.window.strip():
            issues.append(Issue("WARN", "MISSING_FIELD", f"service_levels:{sl.id}.window", "Missing measurement window"))

    # Unused service levels (defined but not referenced by concerns)
    referenced_sl: set[str] = {c.measurement.service_level_id for c in concerns if c.measurement.service_level_id}
    for sl in service_levels:
        if sl.id not in referenced_sl:
            issues.append(Issue("INFO", "UNUSED", f"service_levels:{sl.id}",
                               "Service level is defined but not referenced by any concern.measurement.service_level_id"))

    # ---------------------------
    # Risk validations / gap analysis (MODAF-friendly links)
    # ---------------------------

    for r in risks:
        # owner must exist
        if r.owner and r.owner not in stk_ids:
            issues.append(Issue("ERROR", "BROKEN_LINK", f"risks:{r.id}.owner",
                               f"Risk owner references unknown stakeholder id '{r.owner}'"))

        # affected links must exist
        for cid in r.affected_concerns:
            if cid not in concern_ids:
                issues.append(Issue("ERROR", "BROKEN_LINK", f"risks:{r.id}.affected_concerns",
                                   f"Risk references unknown concern id '{cid}'"))
        for capid in r.affected_capabilities:
            if capid not in cap_ids:
                issues.append(Issue("ERROR", "BROKEN_LINK", f"risks:{r.id}.affected_capabilities",
                                   f"Risk references unknown capability id '{capid}'"))
        for slid in r.threatened_service_levels:
            if slid not in sl_ids:
                issues.append(Issue("ERROR", "BROKEN_LINK", f"risks:{r.id}.threatened_service_levels",
                                   f"Risk references unknown service level id '{slid}'"))

        # recommended coverage: each risk should link at least one concern and one capability
        if not r.affected_concerns:
            issues.append(Issue("WARN", "GAP", f"risks:{r.id}.affected_concerns",
                               "Risk has no affected concerns linked (recommend link to ISO42010 concerns)"))
        if not r.affected_capabilities:
            issues.append(Issue("WARN", "GAP", f"risks:{r.id}.affected_capabilities",
                               "Risk has no affected capabilities linked (recommend link to MODAF capabilities/views)"))

        # mitigation should exist (otherwise TODO-worthy)
        if not r.mitigation.strip():
            issues.append(Issue("WARN", "MISSING_FIELD", f"risks:{r.id}.mitigation",
                               "Risk mitigation is missing/empty"))

        # programmatic risks should reference AcV-2 in linked_views (MODAF acquisition timeline)
        if is_programmatic_risk(r.type):
            if not any("AcV-2" in v for v in r.linked_views):
                issues.append(Issue("WARN", "MISSING_VIEW_LINK", f"risks:{r.id}.linked_views",
                                   "Programmatic/acquisition/schedule risk should link to AcV-2 (Programme Timelines)"))

        # at least one linked view is generally recommended
        if not r.linked_views:
            issues.append(Issue("INFO", "GAP", f"risks:{r.id}.linked_views",
                               "Risk has no linked_views (recommend AV-1 + relevant OV/SV/AcV references)"))

    # Stable ordering
    issues.sort(key=lambda i: (SEVERITY_ORDER.get(i.severity, 99), i.code, i.location, i.message))
    return issues


# ---------------------------
# Rendering (codified AD)
# ---------------------------

DEFAULT_TEMPLATE = r"""
# Architecture Document (Codified): {{ system_name }}

**Document ID:** {{ document_id }}  
**Version:** {{ version }}  
**Date:** {{ date }}  
**Status:** {{ status }}

## 0. Scope
{{ scope | default(todo(), true) }}

---

## 1. Glossary
{% if glossary %}
{% for term in glossary %}
- **{{ term.name }}** — {{ term.description }}
{% endfor %}
{% else %}
- **Stakeholder** — сторона, чьи интересы затрагиваются архитектурой.
- **Concern** — значимый интерес/вопрос, который архитектура должна адресовать.
- **Capability** — способность системы обеспечивать определённый результат (что система должна уметь).
- **SLO (Service Level Objective)** — целевой уровень сервиса (внутренний/операционный).
- **SLA (Service Level Agreement)** — договорённый/контрактный уровень сервиса.
- **Risk** — неопределённость, которая может повлиять на concerns/capabilities/SLA и цели проекта.
{% endif %}

---

## 2. Stakeholders
| ID | Name | Description |
|---|---|---|
{% for s in stakeholders -%}
| {{ s.id }} | {{ s.name | replace("|","\\|") }} | {{ s.description | default("", true) | replace("|","\\|") }} |
{% endfor %}

---

## 3. Concern Registry
> Каждый concern описан один раз и имеет один ID; категории задаются тегами.

| ID | Name | Description | Stakeholders | Tags | Measurement (SLI / SLO / SLA / SL ref) |
|---|---|---|---|---|---|
{% for c in concerns -%}
| {{ c.id }} | {{ c.name | replace("|","\\|") }} | {{ c.description_one_line | replace("|","\\|") }} | {{ c.stakeholders_str | replace("|","\\|") }} | {{ c.tags_str | replace("|","\\|") }} | {{ c.measurement_str | replace("|","\\|") }} |
{% endfor %}

### 3.1 Concern views by tags (no duplication)
**Business:** {{ concerns_by_tag.Business | default(todo(), true) }}  
**Operational:** {{ concerns_by_tag.Operational | default(todo(), true) }}  
**Security:** {{ concerns_by_tag.Security | default(todo(), true) }}  
**Compliance:** {{ concerns_by_tag.Compliance | default(todo(), true) }}  
**Data:** {{ concerns_by_tag.Data | default(todo(), true) }}

---

## 4. Capabilities
| ID | Name | Description | Addresses Concerns | Tags |
|---|---|---|---|---|
{% for cap in capabilities -%}
| {{ cap.id }} | {{ cap.name | replace("|","\\|") }} | {{ cap.description_one_line | replace("|","\\|") }} | {{ cap.concerns_str | replace("|","\\|") }} | {{ cap.tags_str | replace("|","\\|") }} |
{% endfor %}

---

## 5. Risk Register (AV-1 aligned)
| Risk ID | Title | Type | Status | Owner (STK) | Affected Concerns | Affected Capabilities | Threatened SL | Linked Views | Mitigation |
|---|---|---|---|---|---|---|---|---|---|
{% if risks %}
{% for r in risks -%}
| {{ r.id }} | {{ r.title | replace("|","\\|") }} | {{ r.type | replace("|","\\|") }} | {{ r.status | replace("|","\\|") }} | {{ r.owner_display | replace("|","\\|") }} | {{ r.concerns_str | replace("|","\\|") }} | {{ r.caps_str | replace("|","\\|") }} | {{ r.sls_str | replace("|","\\|") }} | {{ r.views_str | replace("|","\\|") }} | {{ r.mitigation_display | replace("|","\\|") }} |
{% endfor %}
{% else %}
| {{ todo() }} | No risks provided | {{ todo() }} | {{ todo() }} | {{ todo() }} | {{ todo() }} | {{ todo() }} | {{ todo() }} | {{ todo() }} | {{ todo() }} |
{% endif %}

---

## 6. Service Level Catalog
{% if service_levels %}
| ID | Name | SLI definition | Window | Exclusions | Target SLO | Contractual SLA |
|---|---|---|---|---|---|---|
{% for sl in service_levels -%}
| {{ sl.id }} | {{ sl.name | replace("|","\\|") }} | {{ sl.sli_definition_one_line | replace("|","\\|") }} | {{ sl.window | replace("|","\\|") }} | {{ sl.exclusions_one_line | replace("|","\\|") }} | {{ sl.target_slo_display | replace("|","\\|") }} | {{ sl.contractual_sla_display | replace("|","\\|") }} |
{% endfor %}
{% else %}
No service level catalog provided. {{ todo() }}
{% endif %}

---

## 7. Traceability (high level)
### 7.1 Stakeholder → Concerns
{% for s in stakeholders -%}
- **{{ s.id }} {{ s.name }}** → {{ stakeholder_to_concerns.get(s.id) | default(todo(), true) }}
{% endfor %}

### 7.2 Concern → Capabilities
{% for c in concerns_src -%}
- **{{ c.id }} {{ c.name }}** → {{ concern_to_capabilities.get(c.id) | default(todo(), true) }}
{% endfor %}

### 7.3 Risk → (Concerns, Capabilities, SL)
{% if risks_trace %}
{% for line in risks_trace -%}
- {{ line }}
{% endfor %}
{% else %}
- {{ todo() }} No risks to trace.
{% endif %}

---

## 8. Gaps / TODO list (generated)
{% if gaps %}
| Severity | Code | Location | Message |
|---|---|---|---|
{% for g in gaps -%}
| {{ g.severity }} | {{ g.code }} | {{ g.location | replace("|","\\|") }} | {{ g.message | replace("|","\\|") }} |
{% endfor %}
{% else %}
No gaps found.
{% endif %}
"""


def render_ad(
    stakeholders: List[Stakeholder],
    concerns_src: List[Concern],
    capabilities: List[Capability],
    service_levels: List[ServiceLevel],
    risks_src: List[Risk],
    issues: List[Issue],
    meta: Dict[str, Any],
    template_text: str,
) -> str:
    env = Environment(
        loader=BaseLoader(),
        autoescape=False,
        undefined=StrictUndefined,
        trim_blocks=True,
        lstrip_blocks=True,
    )
    env.globals["todo"] = todo

    # Maps for traceability
    stakeholder_to_concerns: Dict[str, List[str]] = {s.id: [] for s in stakeholders}
    for c in concerns_src:
        for sid in c.stakeholders:
            stakeholder_to_concerns.setdefault(sid, []).append(c.id)

    concern_to_capabilities: Dict[str, List[str]] = {c.id: [] for c in concerns_src}
    for cap in capabilities:
        for cid in cap.addresses_concerns:
            concern_to_capabilities.setdefault(cid, []).append(cap.id)

    # Tag views for concerns
    concerns_by_tag: Dict[str, str] = {}
    all_tags = ["Business", "Operational", "Security", "Compliance", "Data"]
    for t in all_tags:
        ids = [c.id for c in concerns_src if t in {x.strip() for x in c.tags}]
        concerns_by_tag[t] = ", ".join(ids) if ids else ""

    def fmt_measurement(c: Concern) -> str:
        m = c.measurement
        sli = m.sli or todo()
        slo = m.slo or todo()
        sla = m.sla or todo()
        slref = m.service_level_id or ""
        slref_display = slref if slref else todo()
        return f"SLI={sli}; SLO={slo}; SLA={sla}; SL={slref_display}"

    # Render rows (with inline TODO)
    concerns_rows = [{
        "id": c.id,
        "name": c.name or todo(),
        "description_one_line": as_one_line(c.description) or todo(),
        "stakeholders_str": ", ".join(c.stakeholders) if c.stakeholders else todo(),
        "tags_str": ", ".join(c.tags) if c.tags else todo(),
        "measurement_str": fmt_measurement(c),
    } for c in concerns_src]

    capabilities_rows = []
    for cap in capabilities:
        cap_tags = cap.tags[:] if cap.tags else keyword_classify_capability(cap)
        capabilities_rows.append({
            "id": cap.id,
            "name": cap.name or todo(),
            "description_one_line": as_one_line(cap.description) or todo(),
            "concerns_str": ", ".join(cap.addresses_concerns) if cap.addresses_concerns else todo(),
            "tags_str": ", ".join(cap_tags) if cap_tags else todo(),
        })

    service_levels_rows = [{
        "id": sl.id,
        "name": sl.name or todo(),
        "sli_definition_one_line": as_one_line(sl.sli_definition) or todo(),
        "window": sl.window or todo(),
        "exclusions_one_line": as_one_line(sl.exclusions) if sl.exclusions else "",
        "target_slo_display": sl.target_slo if sl.target_slo else todo(),
        "contractual_sla_display": sl.contractual_sla if sl.contractual_sla else todo(),
    } for sl in service_levels]

    # Risks rows + risk trace lines
    risks_rows = []
    risks_trace: List[str] = []
    for r in risks_src:
        owner_display = r.owner if r.owner else todo()
        concerns_str = ", ".join(r.affected_concerns) if r.affected_concerns else todo()
        caps_str = ", ".join(r.affected_capabilities) if r.affected_capabilities else todo()
        sls_str = ", ".join(r.threatened_service_levels) if r.threatened_service_levels else ""
        sls_display = sls_str if sls_str else todo()
        views_str = ", ".join(r.linked_views) if r.linked_views else todo()
        mitigation_display = as_one_line(r.mitigation) if r.mitigation.strip() else todo()

        risks_rows.append({
            "id": r.id,
            "title": r.title or todo(),
            "type": r.type or todo(),
            "status": r.status or todo(),
            "owner_display": owner_display,
            "concerns_str": concerns_str,
            "caps_str": caps_str,
            "sls_str": sls_display,
            "views_str": views_str,
            "mitigation_display": mitigation_display,
        })

        risks_trace.append(
            f"**{r.id} {r.title}** → Concerns: {concerns_str}; Capabilities: {caps_str}; SL: {sls_display}"
        )

    stakeholder_to_concerns_str = {k: ", ".join(sorted(v)) for k, v in stakeholder_to_concerns.items() if v}
    concern_to_capabilities_str = {k: ", ".join(sorted(v)) for k, v in concern_to_capabilities.items() if v}

    # Show WARN/ERROR in gaps by default
    gaps = [i for i in issues if i.severity in ("ERROR", "WARN")]

    template = env.from_string(template_text)
    return template.render(
        system_name=meta.get("system_name", "RCS") or "RCS",
        document_id=meta.get("document_id", "AD-000") or "AD-000",
        version=meta.get("version", "0.1") or "0.1",
        date=meta.get("date", dt.date.today().isoformat()),
        status=meta.get("status", "Draft") or "Draft",
        scope=meta.get("scope", ""),
        glossary=meta.get("glossary", []),

        stakeholders=stakeholders,
        concerns=concerns_rows,
        concerns_src=concerns_src,
        concerns_by_tag=concerns_by_tag,

        capabilities=capabilities_rows,

        risks=risks_rows,
        risks_trace=risks_trace,

        service_levels=service_levels_rows,

        stakeholder_to_concerns=stakeholder_to_concerns_str,
        concern_to_capabilities=concern_to_capabilities_str,

        gaps=gaps,
    )


# ---------------------------
# Reports
# ---------------------------

def write_text(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text, encoding="utf-8")


def write_json(path: Path, obj: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(obj, ensure_ascii=False, indent=2), encoding="utf-8")


def issues_to_gaps_md(issues: List[Issue]) -> str:
    gaps = [i for i in issues if i.severity in ("ERROR", "WARN")]
    if not gaps:
        return "# Gaps / TODO\n\nNo gaps found.\n"

    lines = [
        "# Gaps / TODO",
        "",
        "| Severity | Code | Location | Message |",
        "|---|---|---|---|",
    ]
    for g in gaps:
        loc = g.location.replace("|", "\\|")
        msg = g.message.replace("|", "\\|")
        lines.append(f"| {g.severity} | {g.code} | {loc} | {msg} |")
    lines.append("")
    return "\n".join(lines)




def markdown_to_docx(markdown_text: str, out_path: Path) -> None:
    try:
        from docx import Document  # type: ignore
    except Exception as exc:  # noqa: BLE001
        raise RuntimeError("python-docx is required for DOCX output. Install: pip install python-docx") from exc

    doc = Document()
    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip()
        if not line:
            doc.add_paragraph("")
            continue

        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            continue
        if line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
            continue
        if re.match(r'^\d+\.\s+', line):
            content = re.sub(r'^\d+\.\s+', '', line, count=1)
            doc.add_paragraph(content.strip(), style="List Number")
            continue

        doc.add_paragraph(line)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))

def summarize_issues(issues: List[Issue]) -> Dict[str, int]:
    out = {"ERROR": 0, "WARN": 0, "INFO": 0}
    for i in issues:
        out[i.severity] = out.get(i.severity, 0) + 1
    return out


# ---------------------------
# Build / CLI
# ---------------------------

def build_all(spec_dir: Path) -> Tuple[List[Stakeholder], List[Concern], List[Capability], List[ServiceLevel], List[Risk], List[Issue]]:
    loaded, load_issues = load_spec(spec_dir)
    issues: List[Issue] = list(load_issues)

    stakeholders = parse_stakeholders(loaded.get("stakeholders", {}), issues)
    concerns = parse_concerns(loaded.get("concerns", {}), issues)
    capabilities = parse_capabilities(loaded.get("capabilities", {}), issues)
    service_levels = parse_service_levels(loaded.get("service_levels", {}), issues)
    risks = parse_risks(loaded.get("risks", {}), issues)

    issues.extend(analyze(stakeholders, concerns, capabilities, service_levels, risks))
    issues.sort(key=lambda i: (SEVERITY_ORDER.get(i.severity, 99), i.code, i.location, i.message))

    return stakeholders, concerns, capabilities, service_levels, risks, issues


def cmd_validate(args: argparse.Namespace) -> int:
    spec_dir = Path(args.spec_dir)
    _, _, _, _, _, issues = build_all(spec_dir)

    report = {
        "generated_at": now_iso(),
        "spec_dir": str(spec_dir.resolve()),
        "summary": summarize_issues(issues),
        "issues": [i.as_dict() for i in issues],
    }

    if args.report:
        write_json(Path(args.report), report)

    summary = report["summary"]
    print(f"Validation summary: ERROR={summary.get('ERROR',0)} WARN={summary.get('WARN',0)} INFO={summary.get('INFO',0)}")

    if summary.get("ERROR", 0) > 0:
        return 2
    if args.fail_on_warn and summary.get("WARN", 0) > 0:
        return 1
    return 0


def cmd_build(args: argparse.Namespace) -> int:
    spec_dir = Path(args.spec_dir)
    stakeholders, concerns, capabilities, service_levels, risks, issues = build_all(spec_dir)

    summary = summarize_issues(issues)
    report = {
        "generated_at": now_iso(),
        "spec_dir": str(spec_dir.resolve()),
        "summary": summary,
        "issues": [i.as_dict() for i in issues],
    }

    out_path = Path(args.out)
    gaps_path = Path(args.gaps) if args.gaps else out_path.with_name("gaps.md")
    report_path = Path(args.report) if args.report else out_path.with_name("validation_report.json")

    template_text = DEFAULT_TEMPLATE
    if args.template:
        template_text = Path(args.template).read_text(encoding="utf-8")

    meta = {
        "system_name": args.system_name,
        "document_id": args.document_id,
        "version": args.version,
        "date": args.date or dt.date.today().isoformat(),
        "status": args.status,
        "scope": args.scope,
        "glossary": [],
    }

    ad_text = render_ad(
        stakeholders=stakeholders,
        concerns_src=concerns,
        capabilities=capabilities,
        service_levels=service_levels,
        risks_src=risks,
        issues=issues,
        meta=meta,
        template_text=template_text,
    )

    out_format = (args.format or "").strip().lower()
    if not out_format:
        out_format = "docx" if out_path.suffix.lower() == ".docx" else "md"

    if out_format not in ("md", "docx"):
        raise ValueError(f"Unsupported format: {out_format}")

    if out_format == "docx":
        markdown_to_docx(ad_text, out_path)
    else:
        write_text(out_path, ad_text)
    write_text(gaps_path, issues_to_gaps_md(issues))
    write_json(report_path, report)

    print(f"Built AD: {out_path}")
    print(f"Gaps: {gaps_path}")
    print(f"Report: {report_path}")
    print(f"Summary: ERROR={summary.get('ERROR',0)} WARN={summary.get('WARN',0)} INFO={summary.get('INFO',0)}")

    if summary.get("ERROR", 0) > 0:
        return 2
    if args.fail_on_warn and summary.get("WARN", 0) > 0:
        return 1
    return 0


def make_parser() -> argparse.ArgumentParser:
    p = argparse.ArgumentParser(prog="adtool.py", description="Deterministic AD builder (codified) + Risk support")
    sub = p.add_subparsers(dest="cmd", required=True)

    pv = sub.add_parser("validate", help="Validate spec + run gap analysis (no rendering)")
    pv.add_argument("spec_dir", help="Directory with spec YAML files")
    pv.add_argument("--report", help="Path to write validation_report.json")
    pv.add_argument("--fail-on-warn", action="store_true", help="Exit non-zero if WARN exists")
    pv.set_defaults(func=cmd_validate)

    pb = sub.add_parser("build", help="Validate + analyze + render AD (codified)")
    pb.add_argument("spec_dir", help="Directory with spec YAML files")
    pb.add_argument("--out", required=True, help="Output AD path (e.g., ad/output/AD_codified.md or .docx)")
    pb.add_argument("--gaps", help="Output gaps markdown path (default: alongside --out as gaps.md)")
    pb.add_argument("--report", help="Output report json path (default: alongside --out as validation_report.json)")
    pb.add_argument("--template", help="Optional Jinja2 markdown template path (overrides built-in template)")
    pb.add_argument("--format", choices=["md", "docx"], default="", help="Output format (default: inferred from --out suffix)")
    pb.add_argument("--fail-on-warn", action="store_true", help="Exit non-zero if WARN exists")

    pb.add_argument("--system-name", default="RCS")
    pb.add_argument("--document-id", default="AD-RCS-001")
    pb.add_argument("--version", default="0.1")
    pb.add_argument("--date", default="")
    pb.add_argument("--status", default="Draft")
    pb.add_argument("--scope", default="")

    pb.set_defaults(func=cmd_build)

    return p


def main() -> int:
    parser = make_parser()
    args = parser.parse_args()
    return int(args.func(args))


if __name__ == "__main__":
    raise SystemExit(main())
